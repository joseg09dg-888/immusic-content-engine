"""Convierte docs/libro/music_business_para_todos.md a un .docx con identidad
de marca IM Music (portada violeta, headings en violeta, tablas, listas)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "libro" / "music_business_para_todos.md"
COVER_IMG = ROOT / "docs" / "libro" / "portada_music_business.png"
OUT_PATH = ROOT / "docs" / "libro" / "Music_Business_Para_Todos_IM_Music.docx"

VIOLETA = RGBColor(0x5E, 0x17, 0xEB)
NEGRO = RGBColor(0x00, 0x00, 0x00)
CREMA = RGBColor(0xF2, 0xED, 0xE5)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTO = RGBColor(0x22, 0x22, 0x22)

HEADING_FONT = "Bahnschrift"
BODY_FONT = "Georgia"


def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_page_break(doc):
    doc.add_page_break()


def build_cover(doc):
    section = doc.sections[0]
    page_w = section.page_width
    page_h = section.page_height
    # Remove margins for this section only so the image bleeds full-page
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(COVER_IMG), width=page_w, height=page_h)

    add_page_break(doc)

    # New section for the rest of the book, with normal margins restored
    from docx.enum.section import WD_SECTION
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.left_margin = Inches(0.7)
    new_section.right_margin = Inches(0.7)
    new_section.top_margin = Inches(0.7)
    new_section.bottom_margin = Inches(0.7)


def style_heading(p, text, size, color, font=HEADING_FONT, space_before=18, space_after=8, align_center=False):
    p.text = ""
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = font
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_inline_runs(paragraph, text, base_size=11):
    """Parse **bold** markdown within a line and add runs."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = VIOLETA
        else:
            r = paragraph.add_run(part)
        r.font.name = BODY_FONT
        r.font.size = Pt(base_size)


def add_table_from_lines(doc, lines):
    rows = [l.strip() for l in lines if l.strip().startswith("|")]
    rows = [r for r in rows if not re.match(r"^\|[\s\-:|]+\|$", r)]
    data = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
    if not data:
        return
    n_cols = len(data[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(data):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(re.sub(r"\*\*", "", val))
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
            if i == 0:
                r.bold = True
                r.font.color.rgb = BLANCO
                set_cell_background(cells[j], "5E17EB")
            else:
                r.font.color.rgb = GRIS_TEXTO
    doc.add_paragraph()


def build_body(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buffer = []
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buffer))
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.left_indent = Inches(0.3)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F2EDE5")
                p._p.get_or_add_pPr().append(shd)
                doc.add_paragraph()
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            add_page_break(doc)
            p = doc.add_paragraph()
            style_heading(p, line[2:].strip(), 26, VIOLETA, space_before=0, space_after=16)
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            style_heading(p, line[3:].strip(), 18, VIOLETA, space_before=20, space_after=10)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            style_heading(p, line[4:].strip(), 15, NEGRO, space_before=14, space_after=6)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_from_lines(doc, table_lines)
            continue

        if re.match(r"^\s*-\s\[.\]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^\s*-\s\[.\]\s", "☐ ", line)
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r"^\s*-\s", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line.strip()[2:])
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\s*\d+\.\s", "", line)
            add_inline_runs(p, text)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip().strip("*"))
            r.italic = True
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline_runs(p, line.strip())
        i += 1


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(8)

    build_cover(doc)
    build_body(doc, md_text)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OK: {OUT_PATH} ({OUT_PATH.stat().st_size // 1024}KB)")


if __name__ == "__main__":
    main()
