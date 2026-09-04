"""Convierte docs/libro/brief_creativo_video.md a un .docx con identidad de
marca IM Music, para entregar al creador de contenido del video del libro."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "libro" / "brief_creativo_video.md"
OUT_PATH = ROOT / "docs" / "libro" / "Brief_Creativo_Video_Libro_IM_Music.docx"

VIOLETA = RGBColor(0x5E, 0x17, 0xEB)
NEGRO = RGBColor(0x00, 0x00, 0x00)
GRIS_TEXTO = RGBColor(0x22, 0x22, 0x22)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

HEADING_FONT = "Bahnschrift"
BODY_FONT = "Georgia"


def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def style_heading(p, text, size, color, font=HEADING_FONT, space_before=18, space_after=8):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = font
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_inline_runs(paragraph, text, base_size=11):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = VIOLETA
        else:
            r = paragraph.add_run(part)
        r.font.name = BODY_FONT
        r.font.size = Pt(base_size)


def add_table_from_lines(doc, lines):
    rows = [l.strip() for l in lines if l.strip().startswith("|")]
    rows = [r for r in rows if not re.match(r"^\|[\s\-:|]+\|$", r)]
    data = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
    if not data:
        return
    n_cols = len(data[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(data):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            if j >= n_cols:
                break
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            r = p.add_run(re.sub(r"\*\*", "", val))
            r.font.size = Pt(9.5)
            r.font.name = BODY_FONT
            if i == 0:
                r.bold = True
                r.font.color.rgb = BLANCO
                set_cell_background(cells[j], "5E17EB")
            else:
                r.font.color.rgb = GRIS_TEXTO
    doc.add_paragraph()


def build_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IM MUSIC")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = VIOLETA
    r.font.name = HEADING_FONT
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Documento interno — Brief para creador de contenido")
    r2.font.size = Pt(11)
    r2.italic = True
    r2.font.color.rgb = GRIS_TEXTO
    r2.font.name = BODY_FONT
    p2.paragraph_format.space_after = Pt(20)


def build_body(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            style_heading(p, line[2:].strip(), 22, VIOLETA, space_before=0, space_after=14)
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            style_heading(p, line[3:].strip(), 16, VIOLETA, space_before=18, space_after=8)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            style_heading(p, line[4:].strip(), 13, NEGRO, space_before=12, space_after=6)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table_from_lines(doc, table_lines)
            continue

        if re.match(r"^\s*-\s\[.\]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^\s*-\s\[.\]\s", "☐ ", line)
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\s*\d+\.\s", "", line)
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r"^\s*-\s", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line.strip()[2:])
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip().strip("*"))
            r.italic = True
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline_runs(p, line.strip())
        i += 1


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(8)

    build_cover(doc)
    build_body(doc, md_text)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OK: {OUT_PATH} ({OUT_PATH.stat().st_size // 1024}KB)")


if __name__ == "__main__":
    main()
