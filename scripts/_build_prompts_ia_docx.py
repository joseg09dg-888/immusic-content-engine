"""Convierte docs/libro/prompts_creador_ia_video.md a un .docx con identidad
de marca IM Music — prompts maestros para crear la creadora de contenido IA."""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "docs" / "libro" / "prompts_creador_ia_video.md"
OUT_PATH = ROOT / "docs" / "libro" / "Prompts_Creadora_IA_Video_Libro_IM_Music.docx"

VIOLETA = RGBColor(0x5E, 0x17, 0xEB)
NEGRO = RGBColor(0x00, 0x00, 0x00)
GRIS_TEXTO = RGBColor(0x22, 0x22, 0x22)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

HEADING_FONT = "Bahnschrift"
BODY_FONT = "Georgia"


def style_heading(p, text, size, color, font=HEADING_FONT, space_before=18, space_after=8):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = font
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_inline_runs(paragraph, text, base_size=11):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = VIOLETA
        else:
            r = paragraph.add_run(part)
        r.font.name = BODY_FONT
        r.font.size = Pt(base_size)


def build_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IM MUSIC")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = VIOLETA
    r.font.name = HEADING_FONT
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Documento interno — Prompts maestros para creadora de contenido IA")
    r2.font.size = Pt(11)
    r2.italic = True
    r2.font.color.rgb = GRIS_TEXTO
    r2.font.name = BODY_FONT
    p2.paragraph_format.space_after = Pt(20)


def build_body(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buffer = []
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                in_code = False
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buffer))
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.left_indent = Inches(0.3)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "F2EDE5")
                p._p.get_or_add_pPr().append(shd)
                doc.add_paragraph()
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            style_heading(p, line[2:].strip(), 22, VIOLETA, space_before=0, space_after=14)
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            style_heading(p, line[3:].strip(), 16, VIOLETA, space_before=18, space_after=8)
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_paragraph()
            style_heading(p, line[4:].strip(), 13, NEGRO, space_before=12, space_after=6)
            i += 1
            continue

        if re.match(r"^\s*-\s\[.\]\s", line):
            p = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^\s*-\s\[.\]\s", "☐ ", line)
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\s*\d+\.\s", "", line)
            add_inline_runs(p, text)
            i += 1
            continue

        if re.match(r"^\s*-\s", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line.strip()[2:])
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(line.strip().strip("*"))
            r.italic = True
            r.font.size = Pt(10)
            r.font.name = BODY_FONT
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        add_inline_runs(p, line.strip())
        i += 1


def main():
    md_text = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(8)

    build_cover(doc)
    build_body(doc, md_text)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OK: {OUT_PATH} ({OUT_PATH.stat().st_size // 1024}KB)")


if __name__ == "__main__":
    main()
